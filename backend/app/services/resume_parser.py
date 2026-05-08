import asyncio
import io
import re
from typing import Any, Optional


_BULLET_LINE_RE = re.compile(r"^\s*(?:[\u2022\u00b7\u25cf\u25aa\u25e6*+-]|\d+[.)\u3001]|[\uff08(]?\d+[\uff09)])\s+")
_SECTION_HEADINGS = (
    "\u5de5\u4f5c\u7ecf\u5386",
    "\u5b9e\u4e60\u7ecf\u5386",
    "\u9879\u76ee\u7ecf\u5386",
    "\u6559\u80b2\u7ecf\u5386",
    "\u6821\u56ed\u7ecf\u5386",
    "\u4e2a\u4eba\u7ecf\u5386",
    "\u5b9e\u8df5\u7ecf\u5386",
    "\u83b7\u5956\u7ecf\u5386",
    "\u6280\u80fd",
    "\u8bc1\u4e66",
    "\u81ea\u6211\u8bc4\u4ef7",
    "\u4e2a\u4eba\u603b\u7ed3",
)
_SECTION_HEADING_RE = re.compile(
    r"^\s*(?:"
    + "|".join(re.escape(item) for item in _SECTION_HEADINGS)
    + r"|WORK\s+EXPERIENCE|EXPERIENCE|PROJECTS?|EDUCATION|SKILLS?|CERTIFICATIONS?"
    + r")\s*[:\uff1a]?\s*$",
    re.IGNORECASE,
)
_DATE_LINE_RE = re.compile(r"\b(?:19|20)\d{2}(?:[./-]\d{1,2})?\b")
_SPACING_RE = re.compile(r"[ \t\u00a0]+")


def _clean_pdf_text(text: str) -> str:
    text = (text or "").replace("\u00a0", " ").replace("\u200b", "")
    text = _SPACING_RE.sub(" ", text)
    return text.strip()


def _looks_like_hard_break(line: str) -> bool:
    if not line:
        return True
    if _BULLET_LINE_RE.match(line) or _SECTION_HEADING_RE.match(line):
        return True
    if _DATE_LINE_RE.search(line):
        return True
    if re.search(r"(?:\u81f3\u4eca|Present|present)\s*$", line):
        return True
    return False


def _is_sentence_end(text: str) -> bool:
    return text.endswith(("\u3002", "\uff01", "\uff1f", ".", "!", "?", ";", "\uff1b", ":", "\uff1a"))


def _normalize_extracted_text(text: str) -> str:
    """Join PDF visual wraps while keeping real headings and bullet items."""
    lines = [_clean_pdf_text(line) for line in (text or "").replace("\r\n", "\n").replace("\r", "\n").split("\n")]
    blocks: list[str] = []
    current = ""

    for line in lines:
        if not line:
            if current:
                blocks.append(current.strip())
                current = ""
            continue

        if _looks_like_hard_break(line):
            if current:
                blocks.append(current.strip())
            current = line
            continue

        if not current:
            current = line
            continue

        if _SECTION_HEADING_RE.match(current) or _is_sentence_end(current):
            blocks.append(current.strip())
            current = line
        else:
            current = f"{current} {line}"

    if current:
        blocks.append(current.strip())

    return "\n".join(blocks)


def _block_text(block: Any) -> str:
    if isinstance(block, dict):
        lines: list[str] = []
        for line in block.get("lines") or []:
            parts = [_clean_pdf_text(str(span.get("text") or "")) for span in line.get("spans") or []]
            text = _clean_pdf_text("".join(parts))
            if text:
                lines.append(text)
        return "\n".join(lines)
    if isinstance(block, (tuple, list)) and len(block) >= 5:
        return str(block[4] or "")
    return ""


def _block_bbox(block: Any) -> tuple[float, float, float, float]:
    if isinstance(block, dict):
        bbox = block.get("bbox") or (0, 0, 0, 0)
    elif isinstance(block, (tuple, list)) and len(block) >= 4:
        bbox = block[:4]
    else:
        bbox = (0, 0, 0, 0)
    try:
        return float(bbox[0]), float(bbox[1]), float(bbox[2]), float(bbox[3])
    except Exception:
        return 0.0, 0.0, 0.0, 0.0


def _sort_pdf_blocks(blocks: list[Any]) -> list[Any]:
    return sorted(blocks, key=lambda block: (_block_bbox(block)[1], _block_bbox(block)[0]))


def _extract_pdf_with_pymupdf(file_bytes: bytes) -> str:
    import fitz

    doc = fitz.open(stream=file_bytes, filetype="pdf")
    pages: list[str] = []
    try:
        for page in doc:
            try:
                raw_blocks = page.get_text("dict", sort=False).get("blocks", [])
                text_blocks = [block for block in raw_blocks if block.get("type") == 0]
            except Exception:
                text_blocks = page.get_text("blocks", sort=False)

            page_parts: list[str] = []
            for block in _sort_pdf_blocks(text_blocks):
                text = _normalize_extracted_text(_block_text(block))
                if text:
                    page_parts.append(text)
            if page_parts:
                pages.append("\n".join(page_parts))
    finally:
        doc.close()

    return "\n\n".join(pages)


def _extract_pdf_with_pypdf(file_bytes: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(file_bytes))
    texts: list[str] = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            texts.append(_normalize_extracted_text(text))
    return "\n\n".join(item for item in texts if item.strip())


def _parse_pdf_sync(file_bytes: bytes) -> str:
    """Extract resume text from PDF, preserving paragraph-ish blocks where possible."""
    try:
        text = _extract_pdf_with_pymupdf(file_bytes)
    except Exception:
        text = ""

    if text.strip():
        return text

    return _extract_pdf_with_pypdf(file_bytes)


def _parse_docx_sync(file_bytes: bytes) -> str:
    """Extract plain text from Word (.docx)."""
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    texts: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            texts.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                texts.append(" | ".join(cells))
    return "\n".join(texts)


async def parse_resume_file(filename: str, file_bytes: bytes) -> Optional[str]:
    lower = filename.lower()
    if lower.endswith(".pdf"):
        return await asyncio.to_thread(_parse_pdf_sync, file_bytes)
    if lower.endswith(".docx"):
        return await asyncio.to_thread(_parse_docx_sync, file_bytes)
    return None
